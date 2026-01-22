import os
import logging
import queue
import csv 
import pandas as pd
from docx import Document 

from PyQt5.QtWidgets import (QWidget, QVBoxLayout, QHBoxLayout, 
                            QLabel, QLineEdit, QPushButton, QGroupBox, QProgressBar,
                            QFileDialog, QMessageBox, QTableWidget, QTableWidgetItem,
                            QHeaderView, QTextEdit, QTabWidget, QCheckBox, QGridLayout)
from PyQt5.QtCore import Qt, pyqtSignal
from PyQt5.QtGui import QFont

from tools.nomenclatural import NomenclaturalStreetIndexer, ProcessingThread

logging.basicConfig(level=logging.INFO, filename='logs.log', filemode='w')

class ExcelProcessorApp(QWidget):
    """Виджет для обработки номенклатурных индексов (только Excel)"""
    
    progress_updated = pyqtSignal(int)
    finished_processing = pyqtSignal(pd.DataFrame)
    error_occurred = pyqtSignal(str)
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.indexer = NomenclaturalStreetIndexer(500)  # Ваш существующий класс
        self.file_path = None
        self.current_df = None
        self.processing_thread = None
        
        self.init_ui()
        
    def init_ui(self):
        """Инициализация интерфейса только для Excel обработки"""
        layout = QVBoxLayout(self)
        layout.setContentsMargins(10, 10, 10, 10)
        
        title_label = QLabel("Обработка номенклатурных индексов")
        title_label.setFont(QFont("Arial", 16, QFont.Bold))
        title_label.setAlignment(Qt.AlignCenter)
        layout.addWidget(title_label)
        
        tabs = QTabWidget()
        layout.addWidget(tabs)
        
        processing_tab = QWidget()
        tabs.addTab(processing_tab, "Обработка")
        
        self.preview_tab = QWidget()
        tabs.addTab(self.preview_tab, "Предпросмотр")
        
        self.setup_processing_tab(processing_tab)
        self.setup_preview_tab(self.preview_tab)
        
        self.status_label = QLabel("Готов к работе")
        self.status_label.setStyleSheet("color: green; padding: 5px;")
        layout.addWidget(self.status_label)
        
    def setup_processing_tab(self, tab):
        layout = QVBoxLayout(tab)
        
        origin_group = self.create_origin_group()
        layout.addWidget(origin_group)
        
        file_group = self.create_file_group()
        layout.addWidget(file_group)
        
        settings_group = self.create_settings_group()
        layout.addWidget(settings_group)
        
        self.process_btn = QPushButton("Обработать файл")
        self.process_btn.setStyleSheet(self.get_process_button_style())
        self.process_btn.clicked.connect(self.process_file)
        self.process_btn.setEnabled(False)
        layout.addWidget(self.process_btn)
        
        self.progress_bar = QProgressBar()
        self.progress_bar.setVisible(False)
        layout.addWidget(self.progress_bar)
        
        instruction_group = self.create_instruction_group()
        layout.addWidget(instruction_group)
        
    def create_origin_group(self):
        group = QGroupBox("Начало координат")
        layout = QHBoxLayout(group)
        
        layout.addWidget(QLabel("X:"))
        self.x_entry = QLineEdit()
        self.x_entry.setPlaceholderText("Введите координату X")
        layout.addWidget(self.x_entry)
        
        layout.addWidget(QLabel("Y:"))
        self.y_entry = QLineEdit()
        self.y_entry.setPlaceholderText("Введите координату Y")
        layout.addWidget(self.y_entry)
        
        self.set_origin_btn = QPushButton("Установить начало координат")
        self.set_origin_btn.clicked.connect(self.set_origin)
        layout.addWidget(self.set_origin_btn)
        
        return group
    
    def create_file_group(self):
        group = QGroupBox("Работа с файлом")
        layout = QVBoxLayout(group)
        
        self.load_file_btn = QPushButton("Загрузить Excel файл")
        self.load_file_btn.clicked.connect(self.load_file)
        layout.addWidget(self.load_file_btn)
        
        self.file_label = QLabel("Файл не выбран")
        self.file_label.setStyleSheet("color: red;")
        layout.addWidget(self.file_label)
        
        return group
    
    def create_settings_group(self):
        group = QGroupBox("Настройки обработки")
        layout = QHBoxLayout(group)
        
        layout.addWidget(QLabel("Столбец X:"))
        self.x_col_entry = QLineEdit()
        self.x_col_entry.setText("X")
        layout.addWidget(self.x_col_entry)
        
        layout.addWidget(QLabel("Столбец Y:"))
        self.y_col_entry = QLineEdit()
        self.y_col_entry.setText("Y")
        layout.addWidget(self.y_col_entry)

        layout.addWidget(QLabel('Улицы:'))
        self.street_entry = QLineEdit()
        self.street_entry.setText('SEM9')  
        layout.addWidget(self.street_entry)
        
        return group
    
    def create_instruction_group(self):
        group = QGroupBox("Инструкция")
        layout = QVBoxLayout(group)        
        instruction_text = QTextEdit()
        instruction_text.setPlainText("""1. Установите начало координат (X и Y)
2. Загрузите Excel файл с координатами
3. Укажите названия столбцов с координатами (по умолчанию X и Y)
4. Укажите название столбца с улицами (опционально)
5. Нажмите 'Обработать файл'
6. Выберите место для сохранения результата""")
        instruction_text.setReadOnly(True)
        layout.addWidget(instruction_text)
        
        return group
    
    def setup_preview_tab(self, tab):
        layout = QVBoxLayout(tab)
        
        self.preview_label = QLabel("Данные для предпросмотра появятся после обработки файла")
        self.preview_label.setAlignment(Qt.AlignCenter)
        layout.addWidget(self.preview_label)
        
        self.preview_table = QTableWidget()
        self.preview_table.setAlternatingRowColors(True)
        self.preview_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.preview_table.setVisible(False)
        layout.addWidget(self.preview_table)
    
    def get_process_button_style(self):
        return """
            QPushButton {
                background-color: lightblue;
                font-weight: bold;
                font-size: 14px;
                padding: 10px;
            }
            QPushButton:disabled {
                background-color: #cccccc;
            }
        """
    
    def set_origin(self):
        try:
            x = float(self.x_entry.text())
            y = float(self.y_entry.text())
            self.indexer.set_origin(x, y)
            QMessageBox.information(self, "Успех", f"Начало координат установлено: X={x}, Y={y}")
            self.update_status("Начало координат установлено", "green")
        except ValueError:
            QMessageBox.critical(self, "Ошибка", "Введите корректные числовые значения для координат")
    
    def update_status(self, message, color="green"):
        self.status_label.setText(message)
        self.status_label.setStyleSheet(f"color: {color}; padding: 5px;")
    
    def load_file(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "Выберите Excel файл",
            "",
            "Excel files (*.xlsx *.xls);;All files (*.*)"
        )
        
        if file_path:
            self.file_path = file_path
            self.file_label.setText(os.path.basename(file_path))
            self.file_label.setStyleSheet("color: green;")
            self.process_btn.setEnabled(True)
            self.update_status("Файл загружен, готов к обработке", "green")
    
    def process_file(self):
        if not self.file_path:
            QMessageBox.critical(self, "Ошибка", "Сначала загрузите файл")
            return
            
        if self.indexer.origin_x is None or self.indexer.origin_y is None:
            QMessageBox.critical(self, "Ошибка", "Сначала установите начало координат")
            return
        
        x_col = self.x_col_entry.text().strip()
        y_col = self.y_col_entry.text().strip()
        street_col = self.street_entry.text().strip() 
        
        # Создаем поток обработки
        self.processing_thread = ProcessingThread(
            self.indexer, self.file_path, x_col, y_col, street_col
        )
        self.processing_thread.progress_updated.connect(self.update_progress)
        self.processing_thread.finished_processing.connect(self.on_processing_finished)
        self.processing_thread.error_occurred.connect(self.on_processing_error)
        
        self.progress_bar.setVisible(True)
        self.progress_bar.setValue(0)
        self.process_btn.setEnabled(False)
        self.update_status("Обработка файла...", "blue")
        
        self.processing_thread.start()
    
    def update_progress(self, value):
        self.progress_bar.setValue(value)
    
    def on_processing_finished(self, df):
        self.current_df = df
        self.progress_bar.setVisible(False)
        self.process_btn.setEnabled(True)
        self.update_status("Обработка завершена", "green")
        
        output_path, selected_filter = QFileDialog.getSaveFileName(
            self,
            "Сохранить результат как",
            "",
            "Excel files (*.xlsx);;Word files (*.docx)"
        )
        
        if output_path:
            try:
                if selected_filter == "Excel files (*.xlsx)":
                    if not output_path.endswith('.xlsx'):
                        base_path = output_path.rsplit('.', 1)[0] if '.' in output_path else output_path
                        output_path = base_path + '.xlsx'
                elif selected_filter == "Word files (*.docx)":
                    if not output_path.endswith('.docx'):
                        base_path = output_path.rsplit('.', 1)[0] if '.' in output_path else output_path
                        output_path = base_path + '.docx'
                
                if output_path.endswith('.docx'):
                    doc = Document()
                    title = doc.add_heading('Обработанные данные улиц', 0)
                    
                    has_street = 'Форматированная улица' in df.columns
                    has_index = 'Номенклатурный индекс' in df.columns
                    
                    if has_street and has_index:
                        for _, row in df.iterrows():
                            p = doc.add_paragraph()
                            street = str(row['Форматированная улица'])
                            index = str(row['Номенклатурный индекс'])
                            
                            p.add_run(street)
                            p.add_run('\t')
                            p.add_run(index).bold = True
                    else:
                        table = doc.add_table(rows=len(df)+1, cols=len(df.columns))
                        table.style = 'Table Grid'
                        
                        for i, column in enumerate(df.columns):
                            table.cell(0, i).text = str(column)
                        
                        for i, (idx, row) in enumerate(df.iterrows(), 1):
                            for j, value in enumerate(row):
                                table.cell(i, j).text = str(value)
                    
                    doc.save(output_path)
                else:
                    df.to_excel(output_path, index=False)

                QMessageBox.information(self, "Успех", f"Файл сохранен как:\n{output_path}")
                self.update_status("Файл успешно обработан и сохранен", "green")
                self.show_preview(df)
                
            except Exception as e:
                QMessageBox.critical(self, "Ошибка", f"Ошибка при сохранении файла: {str(e)}")
        else:
            self.update_status("Обработка отменена", "orange")
    
    def on_processing_error(self, error_message):
        self.progress_bar.setVisible(False)
        self.process_btn.setEnabled(True)
        QMessageBox.critical(self, "Ошибка", f"Произошла ошибка при обработке: {error_message}")
        self.update_status("Ошибка обработки", "red")
    
    def show_preview(self, df):
        for child in self.children():
            if isinstance(child, QTabWidget):
                child.setCurrentIndex(1)
                break
        
        self.preview_table.setVisible(True)
        self.preview_label.setVisible(False)
        
        preview_rows = min(10, len(df))
        self.preview_table.setRowCount(preview_rows)
        self.preview_table.setColumnCount(len(df.columns))
        self.preview_table.setHorizontalHeaderLabels(df.columns.tolist())
        
        for i in range(preview_rows):
            for j, column in enumerate(df.columns):
                item = QTableWidgetItem(str(df.iloc[i, j]))
                item.setFlags(item.flags() & ~Qt.ItemIsEditable)
                self.preview_table.setItem(i, j, item)

    def cleanup(self):
        if self.processing_thread and self.processing_thread.isRunning():
            self.processing_thread.terminate()
            self.processing_thread.wait()

class CheckAndMatch(QWidget):
    def __init__(self, hmap, parent=None):
        super().__init__(parent)
        self.hmap=hmap
        self.parent_app=parent

        self.init_ui()

    def init_ui(self):
        """Инициализация интерфейса"""
        layout = QVBoxLayout(self)
        layout.setSpacing(10)
        layout.setContentsMargins(10, 10, 10, 10)
        
        # Группа параметров
        params_group = QGroupBox("Параметры считки объектов")
        params_layout = QGridLayout()
        params_layout.setSpacing(8)
        
        row = 0
        # Проверяемый слой
        params_layout.addWidget(QLabel("Проверяемый слой:"), row, 0)
        self.lbl_check = QLabel("Объект не выбран")
        self.lbl_check.setStyleSheet("color: gray;")
        params_layout.addWidget(self.lbl_check, row, 1)
        self.btn_check = QPushButton("Выбрать")
        self.btn_check.clicked.connect(lambda: self.pick_layer('check_layer', self.lbl_check, "Проверяемый слой"))
        params_layout.addWidget(self.btn_check, row, 2)
        row += 1
        
        # Целевой слой
        params_layout.addWidget(QLabel("Целевой слой:"), row, 0)
        self.lbl_target = QLabel("Объект не выбран")
        self.lbl_target.setStyleSheet("color: gray;")
        params_layout.addWidget(self.lbl_target, row, 1)
        self.btn_target = QPushButton("Выбрать")
        self.btn_target.clicked.connect(lambda: self.pick_layer('target_layer', self.lbl_target, "Целевой слой"))
        params_layout.addWidget(self.btn_target, row, 2)
        row += 1
        
        # Семантика проверяемого слоя
        params_layout.addWidget(QLabel("Семантика проверяемого:"), row, 0)
        self.lbl_sem_check = QLabel("Семантика не выбрана")
        self.lbl_sem_check.setStyleSheet("color: gray;")
        params_layout.addWidget(self.lbl_sem_check, row, 1)
        self.btn_sem_check = QPushButton("Выбрать")
        self.btn_sem_check.clicked.connect(lambda: self.pick_semantic('check_sem', 'check_sem_name', self.lbl_sem_check))
        params_layout.addWidget(self.btn_sem_check, row, 2)
        row += 1
        
        # Семантика целевого слоя
        params_layout.addWidget(QLabel("Семантика целевого:"), row, 0)
        self.lbl_sem_target = QLabel("Семантика не выбрана")
        self.lbl_sem_target.setStyleSheet("color: gray;")
        params_layout.addWidget(self.lbl_sem_target, row, 1)
        self.btn_sem_target = QPushButton("Выбрать")
        self.btn_sem_target.clicked.connect(lambda: self.pick_semantic('target_sem', 'target_sem_name', self.lbl_sem_target))
        params_layout.addWidget(self.btn_sem_target, row, 2)
        row += 1
        
        # Максимальное расстояние
        params_layout.addWidget(QLabel("Макс. расстояние (м):"), row, 0)
        self.entry_dist = QLineEdit("500")
        self.entry_dist.setMaximumWidth(150)
        params_layout.addWidget(self.entry_dist, row, 1, 1, 2)
        row += 1
        
        # Чекбоксы
        self.cb_add_semantics = QCheckBox("Добавлять семантику 'Ошибка соответствия'")
        self.cb_add_semantics.setChecked(True)
        self.cb_add_semantics.stateChanged.connect(
            lambda state: self.update_param('add_semantics_enabled', state == Qt.Checked)
        )
        params_layout.addWidget(self.cb_add_semantics, row, 0, 1, 3)
        row += 1
        
        self.cb_nearest_mode = QCheckBox("Режим ближайшие соседи (один к одному)")
        self.cb_nearest_mode.setChecked(False)
        self.cb_nearest_mode.stateChanged.connect(
            lambda state: self.update_param('nearest_neighbor_mode', state == Qt.Checked)
        )
        params_layout.addWidget(self.cb_nearest_mode, row, 0, 1, 3)
        row += 1
        
        self.cb_ignore_dot = QCheckBox("Не учитывать точки в семантике (если есть '.', то не ошибка)")
        self.cb_ignore_dot.setChecked(False)
        self.cb_ignore_dot.stateChanged.connect(
            lambda state: self.update_param('ignore_dot_semantics', state == Qt.Checked)
        )
        params_layout.addWidget(self.cb_ignore_dot, row, 0, 1, 3)
        row += 1
        
        params_group.setLayout(params_layout)
        layout.addWidget(params_group)
        
        # Группа управления
        control_group = QGroupBox("Управление считкой")
        control_layout = QHBoxLayout()
        
        self.btn_run = QPushButton("Выполнить считку")
        self.btn_run.setStyleSheet("""
            QPushButton {
                background-color: lightblue;
                font-weight: bold;
                padding: 8px;
            }
            QPushButton:disabled {
                background-color: #cccccc;
            }
        """)
        self.btn_run.clicked.connect(self.run_processing)
        control_layout.addWidget(self.btn_run)
        
        self.btn_save = QPushButton("Сохранить отчёт в CSV")
        self.btn_save.setEnabled(False)
        self.btn_save.clicked.connect(self.save_report)
        control_layout.addWidget(self.btn_save)
        
        control_group.setLayout(control_layout)
        layout.addWidget(control_group)
        
        # Прогресс-бар
        self.progress_bar = QProgressBar()
        self.progress_bar.setVisible(False)
        layout.addWidget(self.progress_bar)
        
        # Статус
        self.status_label = QLabel("Готов к работе")
        self.status_label.setStyleSheet("color: green; font-weight: bold;")
        self.status_label.setAlignment(Qt.AlignCenter)
        layout.addWidget(self.status_label)
        
        # Информационная панель
        info_group = QGroupBox("Информация")
        info_layout = QVBoxLayout()
        
        self.info_text = QTextEdit()
        self.info_text.setReadOnly(True)
        self.info_text.setMaximumHeight(100)
        info_layout.addWidget(self.info_text)
        
        info_group.setLayout(info_layout)
        layout.addWidget(info_group)
        
        # Добавляем растягивающий элемент
        layout.addStretch()
        
    def update_param(self, param_name, value):
        """Обновление параметра в логике"""
        self.logic.params[param_name] = value
        
    def pick_layer(self, param_name, label, title):
        """Выбор слоя из карты"""
        try:
            # Вызов MapAPI для выбора слоя
            # temp = mapapi.mapCreateObject(self.hmap)
            # ... вызов диалога ...
            # layer_key = полученный ключ
            
            # Заглушка для примера
            layer_key = f"{title}_key_{param_name}"
            
            self.logic.params[param_name] = layer_key
            label.setText(layer_key)
            label.setStyleSheet("color: black;")
            self.info_text.append(f"Выбран слой: {layer_key}")
            
        except Exception as e:
            self.logic.message_queue.put(("error", f"Ошибка выбора слоя: {str(e)}"))
    def validate_inputs(self):
        """Проверка входных данных"""
        if not self.logic.params['check_layer']:
            return False, "Выберите проверяемый слой"
        if not self.logic.params['target_layer']:
            return False, "Выберите целевой слой"
        if not self.logic.params['check_sem']:
            return False, "Выберите семантику проверяемого слоя"
        if not self.logic.params['target_sem']:
            return False, "Выберите семантику целевого слоя"
        
        try:
            dist = float(self.entry_dist.get())
            if dist <= 0:
                return False, "Расстояние должно быть > 0"
            self.logic.params['max_dist'] = dist
        except ValueError:
            return False, "Введите корректное число для расстояния"
            
        return True, ""
        
    def run_processing(self):
        """Запуск обработки"""
        is_valid, error_msg = self.validate_inputs()
        if not is_valid:
            QMessageBox.warning(self, "Ошибка", error_msg)
            return
            
        # Формируем сообщение подтверждения
        confirm_msg = self.get_confirmation_message()
        reply = QMessageBox.question(self, "Подтверждение", confirm_msg,
                                    QMessageBox.Yes | QMessageBox.No)
        
        if reply == QMessageBox.Yes:
            self.btn_run.setEnabled(False)
            self.btn_save.setEnabled(False)
            self.status_label.setText("Выполняется считка...")
            self.status_label.setStyleSheet("color: blue;")
            self.progress_bar.setVisible(True)
            self.progress_bar.setRange(0, 0)  # Индикатор без конца
            
            # # Запускаем в фоновом потоке
            # self.processing_thread = SemanticsCheckWorker(self.logic, self.hmap)
            # self.processing_thread.start()
            
    def get_confirmation_message(self):
        """Получить сообщение для подтверждения"""
        nearest_mode_desc = "(связь один к одному)" if self.logic.params['nearest_neighbor_mode'] else "(связь один ко многим)"
        
        return (
            f"Проверяемый слой: {self.logic.params['check_layer']}\n"
            f"Целевой слой: {self.logic.params['target_layer']}\n\n"
            f"Сравнивать семантики:\n"
            f" • {self.logic.params['check_sem_name']} (проверяемый)\n"
            f" • {self.logic.params['target_sem_name']} (целевой)\n\n"
            f"Режим ближайшие соседи: {'ВКЛЮЧЕН' if self.logic.params['nearest_neighbor_mode'] else 'ОТКЛЮЧЕН'} {nearest_mode_desc}\n"
            f"Не учитывать точки в семантике: {'ВКЛЮЧЕНО' if self.logic.params['ignore_dot_semantics'] else 'ОТКЛЮЧЕНО'}\n"
            f"Максимальное расстояние: {self.logic.params['max_dist']} м\n"
            f"Добавление 'Ошибка соответствия': {'ВКЛЮЧЕНО' if self.logic.params['add_semantics_enabled'] else 'ОТКЛЮЧЕНО'}\n\n"
            f"Выполнить считку?"
        )
        
    def save_report(self):
        """Сохранение отчёта"""
        filename, _ = QFileDialog.getSaveFileName(
            self, "Сохранить отчёт в CSV", "", "CSV файлы (*.csv);;Все файлы (*.*)"
        )
        
        if filename:
            if not filename.endswith('.csv'):
                filename += '.csv'
                
            success, message = self.save_report_to_csv(filename)
            if success:
                QMessageBox.information(self, "Успех", message)
            else:
                QMessageBox.critical(self, "Ошибка", message)
                
    def save_report_to_csv(self, filename):
        """Сохранение отчёта в CSV файл"""
        if not self.logic.params.get('result_ready', False):
            return False, "Сначала выполните считку"
            
        try:
            failed_all = (self.logic.params['failed_no_sem_value'] + 
                         self.logic.params['failed_by_distance'] + 
                         self.logic.params['failed_by_sem'] + 
                         self.logic.params['failed_by_both'])
            
            success_transfers = self.logic.params['success_transfers']
            max_rows = max(len(failed_all), len(success_transfers), 10)
            
            rows = []
            rows.append(["ОТЧЁТ СЧИТКИ ОБЪЕКТОВ НА СООТВЕТСТВИЕ СЕМАНТИКИ"] + [""] * 9)
            rows.append([""] * 10)
            
            row3 = [""] * 10
            row3[0] = "Ошибки соответствия"
            row3[3] = "УСПЕШНЫЕ СОПОСТАВЛЕНИЯ"
            row3[7] = "СТАТИСТИКА СЧИТКИ"
            rows.append(row3)
            
            row4 = [""] * 10
            row4[0] = "Номер объекта"
            row4[1] = "Причина"
            row4[3] = "Номер проверяемого"
            row4[4] = "Номер целевого"
            row4[5] = "Расстояние (м)"
            row4[7] = "Параметр"
            row4[8] = "Значение"
            rows.append(row4)
            
            statistics = [
                ("Проверяемый слой", self.logic.params['check_layer']),
                ("Целевой слой", self.logic.params['target_layer']),
                ("Семантика проверяемого", f"{self.logic.params['check_sem']} - {self.logic.params['check_sem_name']}"),
                ("Семантика целевого", f"{self.logic.params['target_sem']} - {self.logic.params['target_sem_name']}"),
                ("Макс. расстояние (м)", self.logic.params['max_dist']),
                ("Режим ближайшие соседи", "ВКЛЮЧЕН" if self.logic.params['nearest_neighbor_mode'] else "ОТКЛЮЧЕН"),
                ("Добавление семантики", "ВКЛЮЧЕНО" if self.logic.params['add_semantics_enabled'] else "ОТКЛЮЧЕНО"),
                ("Не учитывать точки в семантике", "ВКЛЮЧЕНО" if self.logic.params['ignore_dot_semantics'] else "ОТКЛЮЧЕНО"),
                ("Всего проверено объектов", self.logic.params['total']),
                ("Успешно сопоставлено", self.logic.params['success_count']),
                ("Не прошли", self.logic.params['total'] - self.logic.params['success_count']),
                ("Отсутствует значение семантики", len(self.logic.params['failed_no_sem_value'])),
                ("Нет объектов в радиусе", len(self.logic.params['failed_by_distance'])),
                ("Несоответствие семантики", len(self.logic.params['failed_by_sem'])),
                ("Нет объектов в радиусе + несоответствие семантики", len(self.logic.params['failed_by_both'])),
            ]
            
            for i in range(max_rows):
                row = [""] * 10
                
                if i < len(failed_all):
                    key, reason = failed_all[i]
                    row[0] = key
                    row[1] = reason
                
                if i < len(success_transfers):
                    check_key, target_key, dist = success_transfers[i]
                    row[3] = check_key
                    row[4] = target_key
                    row[5] = f"{dist:.2f}"
                
                if i < len(statistics):
                    param, value = statistics[i]
                    row[7] = param
                    row[8] = str(value)
                
                rows.append(row)
            
            with open(filename, 'w', newline='', encoding='utf-8-sig') as f:
                writer = csv.writer(f, delimiter=';')
                writer.writerows(rows)
                
            return True, f"Отчёт сохранён в CSV:\n{filename}"
            
        except Exception as e:
            return False, f"Не удалось сохранить CSV:\n{e}"
            
    def check_messages(self):
        """Проверка очереди сообщений"""
        try:
            while True:
                msg_type, msg_data = self.logic.message_queue.get_nowait()
                
                if msg_type == "log":
                    self.info_text.append(f"[ЛОГ] {msg_data}")
                elif msg_type == "error":
                    QMessageBox.critical(self, "Ошибка", msg_data)
                    self.reset_ui_state()
                elif msg_type == "done":
                    self.on_processing_done()
                    
        except queue.Empty:
            pass
            
    def on_processing_done(self):
        """Обработка завершения"""
        total_failed = (len(self.logic.params['failed_by_distance']) + 
                       len(self.logic.params['failed_by_sem']) + 
                       len(self.logic.params['failed_by_both']) + 
                       len(self.logic.params['failed_no_sem_value']))
        
        result_msg = (
            f"Выполнено!\n\n"
            f"Всего проверено: {self.logic.params['total']}\n"
            f"Успешно: {self.logic.params['success_count']}\n"
            f"Неудачно: {total_failed}"
        )
        
        QMessageBox.information(self, "Результат считки", result_msg)
        
        self.btn_save.setEnabled(True)
        self.status_label.setText("Считка завершена")
        self.status_label.setStyleSheet("color: green;")
        self.reset_ui_state()
        
    def reset_ui_state(self):
        """Сброс состояния UI"""
        self.btn_run.setEnabled(True)
        self.progress_bar.setVisible(False)
        
    def cleanup(self):
        """Очистка ресурсов при закрытии"""
        if self.processing_thread and self.processing_thread.is_alive():
            self.processing_thread.join(timeout=1)
        self.timer.stop()            